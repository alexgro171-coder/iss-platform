from django.utils.dateparse import parse_date
from django.http import HttpResponse
from django.db import models
from rest_framework import viewsets, permissions, status
from rest_framework.decorators import api_view, permission_classes, action
from rest_framework.response import Response
from io import BytesIO
import openpyxl
from openpyxl import Workbook

from .models import (
    Client, Worker, UserProfile, UserRole, ActivityLog, LogType, LogAction,
    WorkerDocument, CodCOR, TemplateDocument, GeneratedDocument, TemplateType
)
from .serializers import (
    ClientSerializer, WorkerSerializer, CurrentUserSerializer,
    WorkerDocumentSerializer, CodCORSerializer, TemplateDocumentSerializer,
    GeneratedDocumentSerializer, GenerateDocumentRequestSerializer
)


@api_view(["GET"])
@permission_classes([permissions.IsAuthenticated])
def current_user(request):
    """
    Returnează informațiile utilizatorului curent.
    Endpoint: GET /api/me/
    
    Folosit de frontend pentru a afișa datele utilizatorului logat.
    """
    serializer = CurrentUserSerializer(request.user)
    return Response(serializer.data)


class IsManagementOrReadOnly(permissions.BasePermission):
    """
    Permite full access Management/Admin.
    Ceilalți pot doar citi (GET).
    """

    def has_permission(self, request, view):
        if not request.user.is_authenticated:
            return False

        if request.method in permissions.SAFE_METHODS:
            return True

        try:
            role = request.user.profile.role
        except UserProfile.DoesNotExist:
            role = None

        return role in (UserRole.MANAGEMENT, UserRole.ADMIN)


class AgentCannotDelete(permissions.BasePermission):
    """
    Interzice ștergerea (DELETE) pentru utilizatorii cu rol Agent.
    
    Conform specificațiilor:
    - Agent: "Nu poate șterge înregistrări"
    
    Permite toate celelalte operațiuni (GET, POST, PUT, PATCH).
    """

    def has_permission(self, request, view):
        # Utilizatorul trebuie să fie autentificat
        if not request.user.is_authenticated:
            return False

        # Dacă NU este o cerere DELETE, permitem
        if request.method != "DELETE":
            return True

        # Este DELETE - verificăm rolul
        try:
            role = request.user.profile.role
        except UserProfile.DoesNotExist:
            role = None

        # Agentul NU poate șterge
        if role == UserRole.AGENT:
            return False

        # Expert, Management, Admin pot șterge
        return True


class ClientViewSet(viewsets.ModelViewSet):
    """
    CRUD pentru clienți.
    Scrierea permisă doar Management/Admin (prin IsManagementOrReadOnly).
    """

    queryset = Client.objects.all().order_by("denumire")
    serializer_class = ClientSerializer
    permission_classes = [IsManagementOrReadOnly]


class CodCORViewSet(viewsets.ModelViewSet):
    """
    CRUD pentru nomenclatorul Coduri COR.
    - GET (listare): toți utilizatorii autentificați
    - POST/PUT/DELETE: doar Management/Admin
    """
    queryset = CodCOR.objects.all().order_by('cod')
    serializer_class = CodCORSerializer
    permission_classes = [IsManagementOrReadOnly]

    def get_queryset(self):
        """Filtrare opțională după status activ."""
        queryset = CodCOR.objects.all()
        
        # Filtru doar coduri active (implicit pentru dropdown-uri)
        activ = self.request.query_params.get('activ')
        if activ is not None:
            queryset = queryset.filter(activ=activ.lower() == 'true')
        
        # Căutare după cod sau denumire
        search = self.request.query_params.get('search')
        if search:
            queryset = queryset.filter(
                models.Q(cod__icontains=search) |
                models.Q(denumire_ro__icontains=search) |
                models.Q(denumire_en__icontains=search)
            )
        
        return queryset.order_by('cod')


class WorkerViewSet(viewsets.ModelViewSet):
    """
    CRUD pentru lucrători, cu filtrare și reguli de acces:
    - Agent -> vede doar lucrătorii unde agent = user
    - Agent -> NU poate șterge (DELETE interzis)
    - Expert/Management/Admin -> văd tot și pot șterge
    """

    serializer_class = WorkerSerializer
    permission_classes = [AgentCannotDelete]  # Aplică restricția de ștergere

    def get_queryset(self):
        user = self.request.user

        # Baza: toți lucrătorii
        qs = Worker.objects.select_related("client", "agent").all()

        if not user.is_authenticated:
            return Worker.objects.none()

        # determinăm rolul
        try:
            role = user.profile.role
        except UserProfile.DoesNotExist:
            role = None

        if role == UserRole.AGENT:
            # Agentul vede DOAR lucrătorii introduși de el
            qs = qs.filter(agent=user)
        else:
            # Expert, Management, Admin -> văd tot
            pass

        # ---- Filtre după query params ----
        params = self.request.query_params

        status_val = params.get("status")
        if status_val:
            qs = qs.filter(status=status_val)

        pasaport = params.get("pasaport_nr")
        if pasaport:
            qs = qs.filter(pasaport_nr__icontains=pasaport)

        cetatenie = params.get("cetatenie")
        if cetatenie:
            qs = qs.filter(cetatenie__iexact=cetatenie)

        client_id = params.get("client_id")
        if client_id:
            qs = qs.filter(client_id=client_id)

        # Filtru după cod COR (cod ocupațional)
        cod_cor = params.get("cod_cor")
        if cod_cor:
            qs = qs.filter(cod_cor__icontains=cod_cor)

        # Filtru după județ WP (Work Permit)
        judet_wp = params.get("judet_wp")
        if judet_wp:
            qs = qs.filter(judet_wp__iexact=judet_wp)

        # Filtru după luna și anul WP (data_programare_wp)
        luna_wp = params.get("luna_wp")
        anul_wp = params.get("anul_wp")
        if luna_wp:
            qs = qs.filter(data_programare_wp__month=int(luna_wp))
        if anul_wp:
            qs = qs.filter(data_programare_wp__year=int(anul_wp))

        # Filtru după luna și anul Viză (data_programare_interviu)
        luna_viza = params.get("luna_viza")
        anul_viza = params.get("anul_viza")
        if luna_viza:
            qs = qs.filter(data_programare_interviu__month=int(luna_viza))
        if anul_viza:
            qs = qs.filter(data_programare_interviu__year=int(anul_viza))

        # interval data_introducere
        data_start = params.get("data_start")
        data_end = params.get("data_end")

        if data_start:
            d_start = parse_date(data_start)
            if d_start:
                qs = qs.filter(data_introducere__date__gte=d_start)

        if data_end:
            d_end = parse_date(data_end)
            if d_end:
                qs = qs.filter(data_introducere__date__lte=d_end)

        return qs.order_by("-data_introducere")

    def perform_create(self, serializer):
        """
        Setează automat agentul la utilizatorul curent când se creează un lucrător.
        Aceasta asigură că agenții pot vedea lucrătorii pe care i-au introdus.
        """
        serializer.save(agent=self.request.user)

    @action(detail=False, methods=['get'], url_path='statistics')
    def statistics(self, request):
        """
        Returnează statistici despre lucrători.
        GET /api/workers/statistics/
        Accesibil pentru Expert, Management, Admin.
        """
        # Verifică dacă user-ul are acces (Expert sau mai sus)
        try:
            role = request.user.profile.role
        except UserProfile.DoesNotExist:
            role = None
        
        if role not in [UserRole.EXPERT, UserRole.MANAGEMENT, UserRole.ADMIN]:
            return Response(
                {'detail': 'Nu aveți permisiunea de a accesa statisticile.'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Baza: toți lucrătorii (Expert/Management/Admin văd tot)
        qs = Worker.objects.all()
        
        # Aplicăm filtre din query params
        params = request.query_params
        
        status_filter = params.get('status')
        if status_filter:
            qs = qs.filter(status=status_filter)
        
        cetatenie = params.get('cetatenie')
        if cetatenie:
            qs = qs.filter(cetatenie__iexact=cetatenie)
        
        # Total
        total = qs.count()
        
        # Pe status
        from django.db.models import Count
        by_status = list(
            qs.values('status')
            .annotate(count=Count('id'))
            .order_by('-count')
        )
        
        # Pe țară/cetățenie
        by_country = list(
            qs.values('cetatenie')
            .annotate(count=Count('id'))
            .order_by('-count')[:10]
        )
        
        # Pe client
        by_client = list(
            qs.filter(client__isnull=False)
            .values('client__denumire')
            .annotate(count=Count('id'))
            .order_by('-count')[:10]
        )
        
        # Liste pentru dropdown-uri de filtre
        available_statuses = list(
            Worker.objects.values_list('status', flat=True).distinct().order_by('status')
        )
        available_countries = list(
            Worker.objects.exclude(cetatenie='')
            .values_list('cetatenie', flat=True)
            .distinct()
            .order_by('cetatenie')
        )
        
        return Response({
            'total': total,
            'by_status': by_status,
            'by_country': by_country,
            'by_client': by_client,
            'available_statuses': available_statuses,
            'available_countries': available_countries,
        })

    @action(detail=False, methods=['get'], url_path='bulk-template')
    def bulk_template(self, request):
        """
        Descarcă template Excel pentru import bulk.
        GET /api/workers/bulk-template/
        """
        # Creăm workbook-ul Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Lucrători"

        # Header-uri (coloanele din template)
        headers = [
            'nume', 'prenume', 'pasaport_nr', 'cetatenie', 'stare_civila',
            'copii_intretinere', 'sex', 'data_nasterii', 'oras_domiciliu',
            'data_emitere_pass', 'data_exp_pass', 'autoritate_emitenta_pasaport',
            'dosar_wp_nr', 'data_solicitare_wp', 'data_programare_wp', 'judet_wp', 'cod_cor', 'functie',
            'data_solicitare_viza', 'data_programare_interviu', 'status',
            'cnp', 'data_intrare_ro', 'cim_nr', 'data_emitere_cim',
            'data_depunere_ps', 'data_programare_ps', 'data_emitere_ps',
            'data_expirare_ps', 'adresa_ro', 'client_denumire', 'observatii'
        ]

        # Scriem header-urile
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")

        # Adăugăm un rând exemplu
        example_row = [
            'Popescu', 'Ion', 'AB123456', 'Nepal', 'M',
            0, 'M', '1990-01-15', 'Kathmandu',
            '2023-01-01', '2033-01-01', 'Ministerul Afacerilor Interne',
            'WP-001', '2024-01-01', '2024-02-01', 'București', '721401', 'Muncitor necalificat',
            '2024-02-15', '2024-03-01', 'Aviz solicitat',
            '', '', '', '',
            '', '', '', '',
            '', 'Client Exemplu', 'Observații exemplu'
        ]
        for col, value in enumerate(example_row, 1):
            ws.cell(row=2, column=col, value=value)

        # Ajustăm lățimea coloanelor
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 30)
            ws.column_dimensions[column].width = adjusted_width

        # Salvăm în memory buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=template_import_lucratori.xlsx'
        return response

    @action(detail=False, methods=['post'], url_path='bulk-import')
    def bulk_import(self, request):
        """
        Import bulk lucrători din fișier Excel.
        POST /api/workers/bulk-import/
        """
        # Verificăm permisiunile (doar Management/Admin)
        try:
            role = request.user.profile.role
        except UserProfile.DoesNotExist:
            role = None

        if role not in (UserRole.MANAGEMENT, UserRole.ADMIN):
            return Response(
                {'detail': 'Nu aveți permisiunea de a importa lucrători.'},
                status=status.HTTP_403_FORBIDDEN
            )

        # Verificăm dacă avem fișier
        if 'file' not in request.FILES:
            return Response(
                {'detail': 'Te rog încarcă un fișier Excel.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        file = request.FILES['file']

        # Verificăm extensia
        if not file.name.endswith(('.xlsx', '.xls')):
            return Response(
                {'detail': 'Fișierul trebuie să fie în format Excel (.xlsx sau .xls).'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            # Citim fișierul Excel
            wb = openpyxl.load_workbook(file)
            ws = wb.active

            # Obținem header-urile și le normalizăm (lowercase, fără spații)
            raw_headers = [cell.value for cell in ws[1]]
            headers = []
            import re
            for h in raw_headers:
                if h:
                    # Normalizăm: lowercase, strip
                    normalized = str(h).lower().strip()
                    # Eliminăm tot ce e în paranteze (ex: "(m/nm)", "(yyyy-mm-dd)")
                    normalized = re.sub(r'\([^)]*\)', '', normalized)
                    # Eliminăm asteriscuri și alte caractere speciale
                    normalized = normalized.replace('*', '').replace('.', '').replace(',', '').replace(':', '')
                    # Înlocuim spații multiple cu unul singur, apoi cu underscore
                    normalized = re.sub(r'\s+', '_', normalized)
                    # Eliminăm underscore-uri la început și sfârșit
                    normalized = normalized.strip('_')
                    
                    # Mapări pentru variante comune (inclusiv în română cu diacritice)
                    header_map = {
                        # Pașaport - toate variantele posibile
                        'nr_pasaport': 'pasaport_nr',
                        'nr_pașaport': 'pasaport_nr',
                        'numar_pasaport': 'pasaport_nr',
                        'număr_pasaport': 'pasaport_nr',
                        'număr_pașaport': 'pasaport_nr',
                        'pasaport': 'pasaport_nr',
                        'pașaport': 'pasaport_nr',
                        'passport': 'pasaport_nr',
                        'passport_nr': 'pasaport_nr',
                        'passport_number': 'pasaport_nr',
                        'pașaport_nr': 'pasaport_nr',
                        # Date pașaport
                        'data_emitere_pașaport': 'data_emitere_pass',
                        'data_emitere_pasaport': 'data_emitere_pass',
                        'data_emitere': 'data_emitere_pass',
                        'data_expirare_pașaport': 'data_exp_pass',
                        'data_expirare_pasaport': 'data_exp_pass',
                        'data_expirare': 'data_exp_pass',
                        # Autoritate emitentă pașaport
                        'autoritate_emitenta_pasaport': 'autoritate_emitenta_pasaport',
                        'autoritate_emitentă_pașaport': 'autoritate_emitenta_pasaport',
                        'autoritate_emitenta': 'autoritate_emitenta_pasaport',
                        'autoritate_emitentă': 'autoritate_emitenta_pasaport',
                        'emitent_pasaport': 'autoritate_emitenta_pasaport',
                        'emitent_pașaport': 'autoritate_emitenta_pasaport',
                        'emitent': 'autoritate_emitenta_pasaport',
                        'issuing_authority': 'autoritate_emitenta_pasaport',
                        'passport_authority': 'autoritate_emitenta_pasaport',
                        # Nume/Prenume
                        'first_name': 'prenume',
                        'last_name': 'nume',
                        'family_name': 'nume',
                        'given_name': 'prenume',
                        'name': 'nume',
                        'surname': 'nume',
                        'forename': 'prenume',
                        # Cetățenie
                        'nationality': 'cetatenie',
                        'citizenship': 'cetatenie',
                        'cetățenie': 'cetatenie',
                        'cetăţenie': 'cetatenie',
                        'cetatenie': 'cetatenie',
                        # Date naștere
                        'birth_date': 'data_nasterii',
                        'date_of_birth': 'data_nasterii',
                        'data_nașterii': 'data_nasterii',
                        'data_nastere': 'data_nasterii',
                        # Stare civilă
                        'stare_civilă': 'stare_civila',
                        'stare_civila': 'stare_civila',
                        # Sex
                        'sex': 'sex',
                        'gen': 'sex',
                        # Copii
                        'copii_întreținere': 'copii_intretinere',
                        'copii_intretinere': 'copii_intretinere',
                        'copii_în_întreținere': 'copii_intretinere',
                        'copii': 'copii_intretinere',
                        # Oraș domiciliu
                        'oraș_domiciliu': 'oras_domiciliu',
                        'oras_domiciliu': 'oras_domiciliu',
                        'oraș': 'oras_domiciliu',
                        'oras': 'oras_domiciliu',
                        'domiciliu': 'oras_domiciliu',
                        # Work Permit
                        'județ_wp': 'judet_wp',
                        'judet_wp': 'judet_wp',
                        'județ': 'judet_wp',
                        'judet': 'judet_wp',
                        'dosar_wp_nr': 'dosar_wp_nr',
                        'nr_dosar_wp': 'dosar_wp_nr',
                        'nr_dosar': 'dosar_wp_nr',
                        'data_solicitare_aviz': 'data_solicitare_wp',
                        'data_solicitare_wp': 'data_solicitare_wp',
                        'data_programare_igi': 'data_programare_wp',
                        'data_programare_wp': 'data_programare_wp',
                        # Cod COR
                        'cod_cor': 'cod_cor',
                        'cor': 'cod_cor',
                        # Funcție
                        'functie': 'functie',
                        'funcție': 'functie',
                        'function': 'functie',
                        'job_title': 'functie',
                        'ocupatie': 'functie',
                        'ocupație': 'functie',
                        'post': 'functie',
                        # Viză
                        'data_solicitare_viza': 'data_solicitare_viza',
                        'data_solicitare_viză': 'data_solicitare_viza',
                        'data_programare_interviu': 'data_programare_interviu',
                        'data_interviu': 'data_programare_interviu',
                        # Status
                        'status': 'status',
                        'stare': 'status',
                        # Permis ședere
                        'data_depunere_permis_ședere': 'data_depunere_ps',
                        'data_depunere_ps': 'data_depunere_ps',
                        'data_programare_permis_ședere': 'data_programare_ps',
                        'data_programare_ps': 'data_programare_ps',
                        'data_emitere_permis_ședere': 'data_emitere_ps',
                        'data_emitere_ps': 'data_emitere_ps',
                        'data_expirare_permis_ședere': 'data_expirare_ps',
                        'data_expirare_ps': 'data_expirare_ps',
                        # Date România
                        'cnp': 'cnp',
                        'data_intrare_ro': 'data_intrare_ro',
                        'data_intrare_în_romania': 'data_intrare_ro',
                        'data_intrare_în_ro': 'data_intrare_ro',
                        'cim_nr': 'cim_nr',
                        'nr_cim': 'cim_nr',
                        'data_emitere_cim': 'data_emitere_cim',
                        'adresa_ro': 'adresa_ro',
                        'adresa': 'adresa_ro',
                        'adresă_în_românia': 'adresa_ro',
                        # Client
                        'client': 'client_denumire',
                        'client_denumire': 'client_denumire',
                        'denumire_client': 'client_denumire',
                        # Observații
                        'observații': 'observatii',
                        'observatii': 'observatii',
                        'obs': 'observatii',
                    }
                    headers.append(header_map.get(normalized, normalized))
                else:
                    headers.append(None)
            
            # Debug: returnăm headers detectate dacă nu avem datele corecte
            detected_headers = [h for h in headers if h]
            
            # Tracking pentru coduri COR nou create
            new_cor_codes = []
            
            results = {
                'total': 0,
                'success': 0,
                'errors': 0,
                'details': [],
                'debug_headers': detected_headers[:15],  # Primele 15 headers pentru debug
                'new_cor_codes': [],  # Coduri COR noi adăugate
            }

            # Procesăm fiecare rând (începând de la 2 pentru a sări header-ul)
            for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                # Skip rânduri complet goale
                if not any(row):
                    continue
                    
                results['total'] += 1
                
                # Creăm dict cu datele
                row_data = {}
                for col_idx, value in enumerate(row):
                    if col_idx < len(headers) and headers[col_idx]:
                        # Acceptăm și valori care nu sunt None
                        if value is not None and value != '':
                            row_data[headers[col_idx]] = value
                
                # Debug: la primul rând, afișăm ce date am citit
                if row_idx == 2 and 'debug_first_row' not in results:
                    results['debug_first_row'] = {k: str(v)[:50] for k, v in list(row_data.items())[:10]}

                # Verificăm câmpurile obligatorii
                nume = row_data.get('nume')
                prenume = row_data.get('prenume')
                pasaport = row_data.get('pasaport_nr')
                
                # Skip rânduri fără date obligatorii
                if not nume or not pasaport:
                    # Dacă are alte date dar lipsesc câmpuri obligatorii
                    if any(v for v in row_data.values() if v):
                        results['errors'] += 1
                        missing = []
                        if not nume:
                            missing.append('nume')
                        if not prenume:
                            missing.append('prenume')
                        if not pasaport:
                            missing.append('pasaport_nr')
                        
                        # La primul rând cu eroare, adăugăm info despre headers detectate
                        error_msg = f'Lipsesc câmpuri obligatorii: {", ".join(missing)}'
                        if row_idx == 2 and detected_headers:
                            error_msg += f' (Coloane detectate: {", ".join(detected_headers[:10])}...)'
                        
                        results['details'].append({
                            'row': row_idx,
                            'status': 'error',
                            'message': error_msg
                        })
                    continue

                try:
                    # Verificăm dacă pașaportul există deja
                    pasaport_nr = str(row_data.get('pasaport_nr', '')).strip()
                    if Worker.objects.filter(pasaport_nr=pasaport_nr).exists():
                        results['errors'] += 1
                        results['details'].append({
                            'row': row_idx,
                            'status': 'error',
                            'message': f'Pașaportul {pasaport_nr} există deja în baza de date'
                        })
                        continue

                    # Găsim clientul dacă e specificat
                    client = None
                    client_denumire = row_data.get('client_denumire')
                    if client_denumire:
                        client = Client.objects.filter(denumire__iexact=str(client_denumire).strip()).first()

                    # Procesăm datele
                    def parse_date_value(value):
                        if not value:
                            return None
                        if isinstance(value, str):
                            return parse_date(value)
                        # Dacă e datetime din Excel
                        try:
                            return value.date() if hasattr(value, 'date') else value
                        except:
                            return None

                    def parse_int_value(value):
                        try:
                            return int(value) if value else 0
                        except:
                            return 0

                    # Helper pentru a obține string safe (evită "None")
                    def get_str(key, default=''):
                        val = row_data.get(key)
                        if val is None:
                            return default
                        return str(val).strip()
                    
                    # Verificăm și procesăm Cod COR
                    cod_cor_value = get_str('cod_cor')
                    cod_cor_ref = None
                    if cod_cor_value:
                        # Căutăm codul COR în nomenclator
                        cod_cor_ref = CodCOR.objects.filter(cod=cod_cor_value).first()
                        if not cod_cor_ref:
                            # Codul COR nu există - îl creăm
                            cod_cor_ref = CodCOR.objects.create(
                                cod=cod_cor_value,
                                denumire_ro='[De completat]',
                                denumire_en='[To be completed]',
                                activ=True
                            )
                            # Adăugăm la lista de coduri noi (dacă nu există deja)
                            if cod_cor_value not in new_cor_codes:
                                new_cor_codes.append(cod_cor_value)
                    
                    # Creăm lucrătorul
                    worker = Worker.objects.create(
                        nume=get_str('nume'),
                        prenume=get_str('prenume'),
                        pasaport_nr=pasaport_nr,
                        cetatenie=get_str('cetatenie'),
                        stare_civila=get_str('stare_civila')[:2] if get_str('stare_civila') else '',
                        copii_intretinere=parse_int_value(row_data.get('copii_intretinere')),
                        sex=get_str('sex')[:1].upper() if get_str('sex') else '',
                        data_nasterii=parse_date_value(row_data.get('data_nasterii')),
                        oras_domiciliu=get_str('oras_domiciliu'),
                        data_emitere_pass=parse_date_value(row_data.get('data_emitere_pass')),
                        data_exp_pass=parse_date_value(row_data.get('data_exp_pass')),
                        autoritate_emitenta_pasaport=get_str('autoritate_emitenta_pasaport'),
                        dosar_wp_nr=get_str('dosar_wp_nr'),
                        data_solicitare_wp=parse_date_value(row_data.get('data_solicitare_wp')),
                        data_programare_wp=parse_date_value(row_data.get('data_programare_wp')),
                        judet_wp=get_str('judet_wp'),
                        cod_cor=cod_cor_value,
                        cod_cor_ref=cod_cor_ref,  # Legătură la nomenclatorul CodCOR
                        functie=get_str('functie'),
                        data_solicitare_viza=parse_date_value(row_data.get('data_solicitare_viza')),
                        data_programare_interviu=parse_date_value(row_data.get('data_programare_interviu')),
                        status=get_str('status') or 'Aviz solicitat',
                        cnp=get_str('cnp'),
                        data_intrare_ro=parse_date_value(row_data.get('data_intrare_ro')),
                        cim_nr=get_str('cim_nr'),
                        data_emitere_cim=parse_date_value(row_data.get('data_emitere_cim')),
                        data_depunere_ps=parse_date_value(row_data.get('data_depunere_ps')),
                        data_programare_ps=parse_date_value(row_data.get('data_programare_ps')),
                        data_emitere_ps=parse_date_value(row_data.get('data_emitere_ps')),
                        data_expirare_ps=parse_date_value(row_data.get('data_expirare_ps')),
                        adresa_ro=get_str('adresa_ro'),
                        client=client,
                        agent=request.user,  # Agentul care importă
                        observatii=get_str('observatii'),
                    )

                    results['success'] += 1
                    # Afișăm datele salvate pentru verificare
                    saved_info = f'{worker.nume} {worker.prenume}'
                    if worker.cetatenie:
                        saved_info += f', {worker.cetatenie}'
                    results['details'].append({
                        'row': row_idx,
                        'status': 'success',
                        'message': f'{saved_info} importat cu succes'
                    })

                except Exception as e:
                    results['errors'] += 1
                    results['details'].append({
                        'row': row_idx,
                        'status': 'error',
                        'message': str(e)
                    })

            # Adăugăm codurile COR noi la rezultate
            results['new_cor_codes'] = new_cor_codes
            
            # Logăm importul
            ActivityLog.log(
                log_type=LogType.ACTIVITY,
                action=LogAction.BULK_IMPORT,
                user=request.user,
                details={
                    'message': f'Import bulk: {results["success"]} succes, {results["errors"]} erori',
                    'total': results['total'],
                    'success': results['success'],
                    'errors': results['errors'],
                    'filename': file.name,
                    'new_cor_codes': new_cor_codes,
                },
                request=request
            )

            return Response(results)

        except Exception as e:
            return Response(
                {'detail': f'Eroare la procesarea fișierului: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )


class WorkerDocumentViewSet(viewsets.ModelViewSet):
    """ViewSet pentru gestionarea documentelor lucrătorilor."""
    queryset = WorkerDocument.objects.all()
    serializer_class = WorkerDocumentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        """Filtrare documente după worker_id dacă e specificat."""
        queryset = WorkerDocument.objects.all()
        worker_id = self.request.query_params.get('worker_id')
        if worker_id:
            queryset = queryset.filter(worker_id=worker_id)
        return queryset

    def create(self, request, *args, **kwargs):
        """Upload document nou."""
        file = request.FILES.get('file')
        if not file:
            return Response(
                {'detail': 'Fișierul este obligatoriu.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        worker_id = request.data.get('worker_id')
        if not worker_id:
            return Response(
                {'detail': 'worker_id este obligatoriu.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            worker = Worker.objects.get(pk=worker_id)
        except Worker.DoesNotExist:
            return Response(
                {'detail': 'Lucrătorul nu a fost găsit.'},
                status=status.HTTP_404_NOT_FOUND
            )

        document = WorkerDocument.objects.create(
            worker=worker,
            document_type=request.data.get('document_type', 'altele'),
            file=file,
            original_filename=file.name,
            description=request.data.get('description', ''),
            uploaded_by=request.user,
            file_size=file.size,
        )

        serializer = self.get_serializer(document)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def destroy(self, request, *args, **kwargs):
        """Șterge un document."""
        instance = self.get_object()
        # Ștergem și fișierul fizic
        if instance.file:
            instance.file.delete(save=False)
        self.perform_destroy(instance)
        return Response(status=status.HTTP_204_NO_CONTENT)


class IsExpertOrAbove(permissions.BasePermission):
    """
    Permite acces doar pentru Expert, Management sau Admin.
    Agentul NU are acces.
    """
    def has_permission(self, request, view):
        if not request.user.is_authenticated:
            return False
        try:
            role = request.user.profile.role
        except UserProfile.DoesNotExist:
            return False
        return role in (UserRole.EXPERT, UserRole.MANAGEMENT, UserRole.ADMIN)


class TemplateDocumentViewSet(viewsets.ModelViewSet):
    """
    ViewSet pentru gestionarea template-urilor de documente.
    Acces: Expert, Management, Admin
    """
    queryset = TemplateDocument.objects.all()
    serializer_class = TemplateDocumentSerializer
    permission_classes = [IsExpertOrAbove]

    def get_queryset(self):
        """Filtrare template-uri."""
        queryset = TemplateDocument.objects.all()
        
        # Filtrare după tip
        template_type = self.request.query_params.get('template_type')
        if template_type:
            queryset = queryset.filter(template_type=template_type)
        
        # Filtrare doar active
        active_only = self.request.query_params.get('active_only', 'true')
        if active_only.lower() == 'true':
            queryset = queryset.filter(is_active=True)
        
        return queryset.select_related('uploaded_by')

    @action(detail=False, methods=['get'], url_path='types')
    def list_types(self, request):
        """
        Listează toate tipurile de template-uri disponibile
        cu status-ul lor (are sau nu template activ).
        """
        types = []
        for value, label in TemplateType.choices:
            active_template = TemplateDocument.objects.filter(
                template_type=value, is_active=True
            ).first()
            types.append({
                'value': value,
                'label': label,
                'has_active_template': active_template is not None,
                'active_template_id': active_template.id if active_template else None,
            })
        return Response(types)

    @action(detail=False, methods=['post'], url_path='upload')
    def upload_template(self, request):
        """
        Upload un nou template.
        Dezactivează automat template-ul anterior de același tip.
        """
        file = request.FILES.get('file')
        if not file:
            return Response(
                {'detail': 'Fișierul este obligatoriu.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validare extensie
        if not file.name.endswith('.docx'):
            return Response(
                {'detail': 'Doar fișiere .docx sunt acceptate.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        template_type = request.data.get('template_type')
        if not template_type:
            return Response(
                {'detail': 'template_type este obligatoriu.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validare tip template
        valid_types = [t[0] for t in TemplateType.choices]
        if template_type not in valid_types:
            return Response(
                {'detail': f'Tip template invalid. Valori acceptate: {valid_types}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Creăm template-ul (save() va dezactiva automat pe celălalt)
        template = TemplateDocument.objects.create(
            template_type=template_type,
            file=file,
            original_filename=file.name,
            is_active=True,
            uploaded_by=request.user,
            description=request.data.get('description', ''),
        )
        
        # Log
        ActivityLog.log(
            log_type=LogType.ACTIVITY,
            action=LogAction.UPLOAD,
            user=request.user,
            target=template,
            details={
                'message': f'Template încărcat: {template.get_template_type_display()}',
                'filename': file.name,
            },
            request=request
        )
        
        serializer = self.get_serializer(template)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=['post'], url_path='generate')
    def generate_document(self, request):
        """
        Generează un document pe baza unui template și datelor unui lucrător.
        Înlocuiește placeholder-ele cu date reale.
        """
        from docx import Document as DocxDocument
        import re
        
        serializer = GenerateDocumentRequestSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        template_type = serializer.validated_data['template_type']
        worker_id = serializer.validated_data['worker_id']
        output_format = serializer.validated_data['output_format']
        
        # Obținem template-ul activ
        template = TemplateDocument.objects.filter(
            template_type=template_type, is_active=True
        ).first()
        
        if not template:
            return Response(
                {'detail': 'Nu există template activ pentru acest tip.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Obținem lucrătorul cu toate relațiile
        worker = Worker.objects.select_related('client', 'cod_cor_ref').get(pk=worker_id)
        
        # Construim maparea placeholder-elor
        placeholder_map = self._build_placeholder_map(worker)
        
        # Încărcăm documentul Word
        try:
            doc = DocxDocument(template.file.path)
        except Exception as e:
            return Response(
                {'detail': f'Eroare la încărcarea template-ului: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        # Înlocuim placeholder-ele în paragrafe
        for paragraph in doc.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, placeholder_map)
        
        # Înlocuim placeholder-ele în tabele
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, placeholder_map)
        
        # Salvăm documentul în buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Salvăm în istoricul documentelor generate
        GeneratedDocument.objects.create(
            template=template,
            template_type=template_type,
            worker=worker,
            worker_name=f"{worker.nume} {worker.prenume}",
            generated_by=request.user,
            generated_by_username=request.user.username,
            output_format=output_format,
        )
        
        # Log
        ActivityLog.log(
            log_type=LogType.ACTIVITY,
            action=LogAction.DOWNLOAD,
            user=request.user,
            target=worker,
            details={
                'message': f'Document generat: {template.get_template_type_display()}',
                'template_id': template.id,
                'output_format': output_format,
            },
            request=request
        )
        
        # Generăm răspunsul
        if output_format == 'pdf':
            # Conversie la PDF
            pdf_buffer = self._convert_to_pdf(buffer)
            if pdf_buffer:
                response = HttpResponse(
                    pdf_buffer.getvalue(),
                    content_type='application/pdf'
                )
                filename = f"{template_type}_{worker.nume}_{worker.prenume}.pdf"
            else:
                # Dacă conversia la PDF eșuează, returnăm Word
                response = HttpResponse(
                    buffer.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                filename = f"{template_type}_{worker.nume}_{worker.prenume}.docx"
        else:
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = f"{template_type}_{worker.nume}_{worker.prenume}.docx"
        
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    def _build_placeholder_map(self, worker):
        """Construiește maparea dintre placeholder-e și valori."""
        
        def format_date(date_val):
            if date_val:
                return date_val.strftime('%d.%m.%Y')
            return ''
        
        # Date lucrător
        placeholder_map = {
            # Date personale
            'nume': worker.nume or '',
            'prenume': worker.prenume or '',
            'nume_complet': f"{worker.nume} {worker.prenume}".strip(),
            'cetatenie': worker.cetatenie or '',
            'stare_civila': worker.stare_civila or '',
            'copii_intretinere': str(worker.copii_intretinere),
            'sex': worker.sex or '',
            'data_nasterii': format_date(worker.data_nasterii),
            'oras_domiciliu': worker.oras_domiciliu or '',
            
            # Pașaport
            'pasaport_nr': worker.pasaport_nr or '',
            'nr_pasaport': worker.pasaport_nr or '',
            'data_emitere_pass': format_date(worker.data_emitere_pass),
            'data_exp_pass': format_date(worker.data_exp_pass),
            'data_expirare_pasaport': format_date(worker.data_exp_pass),
            'autoritate_emitenta_pasaport': worker.autoritate_emitenta_pasaport or '',
            
            # Work Permit
            'dosar_wp_nr': worker.dosar_wp_nr or '',
            'data_solicitare_wp': format_date(worker.data_solicitare_wp),
            'data_programare_wp': format_date(worker.data_programare_wp),
            'judet_wp': worker.judet_wp or '',
            
            # Cod COR
            'cod_cor': worker.cod_cor or '',
            'cod_cor_denumire_ro': worker.cod_cor_ref.denumire_ro if worker.cod_cor_ref else '',
            'cod_cor_denumire_en': worker.cod_cor_ref.denumire_en if worker.cod_cor_ref else '',
            'functie': worker.functie or '',
            
            # Viză
            'data_solicitare_viza': format_date(worker.data_solicitare_viza),
            'data_programare_interviu': format_date(worker.data_programare_interviu),
            
            # Status
            'status': worker.status or '',
            
            # România
            'cnp': worker.cnp or '',
            'data_intrare_ro': format_date(worker.data_intrare_ro),
            'cim_nr': worker.cim_nr or '',
            'data_emitere_cim': format_date(worker.data_emitere_cim),
            'data_depunere_ps': format_date(worker.data_depunere_ps),
            'data_programare_ps': format_date(worker.data_programare_ps),
            'data_emitere_ps': format_date(worker.data_emitere_ps),
            'data_expirare_ps': format_date(worker.data_expirare_ps),
            'adresa_ro': worker.adresa_ro or '',
            
            # Observații
            'observatii': worker.observatii or '',
        }
        
        # Date client
        if worker.client:
            placeholder_map.update({
                'client_denumire': worker.client.denumire or '',
                'client_tara': worker.client.tara or '',
                'client_oras': worker.client.oras or '',
                'client_judet': worker.client.judet or '',
                'client_adresa': worker.client.adresa or '',
                'client_cod_fiscal': worker.client.cod_fiscal or '',
            })
        else:
            placeholder_map.update({
                'client_denumire': '',
                'client_tara': '',
                'client_oras': '',
                'client_judet': '',
                'client_adresa': '',
                'client_cod_fiscal': '',
            })
        
        return placeholder_map

    def _replace_placeholders_in_paragraph(self, paragraph, placeholder_map):
        """Înlocuiește placeholder-ele <field> într-un paragraf."""
        import re
        
        # Combinăm tot textul din runs pentru a detecta placeholder-e
        full_text = ''.join(run.text for run in paragraph.runs)
        
        # Găsim toate placeholder-ele
        placeholders = re.findall(r'<([a-z_]+)>', full_text)
        
        if not placeholders:
            return
        
        # Înlocuim în textul complet
        for ph in placeholders:
            value = placeholder_map.get(ph, '')
            full_text = full_text.replace(f'<{ph}>', value)
        
        # Ștergem runs existente și adăugăm textul nou
        if paragraph.runs:
            # Păstrăm formatarea primului run
            first_run = paragraph.runs[0]
            for run in paragraph.runs[1:]:
                run.text = ''
            first_run.text = full_text

    def _convert_to_pdf(self, docx_buffer):
        """
        Convertește documentul Word la PDF.
        Returnează None dacă conversia eșuează.
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from docx import Document as DocxDocument
            
            # Re-încărcăm documentul din buffer
            docx_buffer.seek(0)
            doc = DocxDocument(docx_buffer)
            
            pdf_buffer = BytesIO()
            c = canvas.Canvas(pdf_buffer, pagesize=A4)
            width, height = A4
            
            y_position = height - 50
            line_height = 14
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text.strip():
                    # Wrap text dacă e prea lung
                    words = text.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + " " + word if current_line else word
                        if len(test_line) * 7 > width - 100:  # Aproximare simplă
                            c.drawString(50, y_position, current_line)
                            y_position -= line_height
                            current_line = word
                        else:
                            current_line = test_line
                    
                    if current_line:
                        c.drawString(50, y_position, current_line)
                        y_position -= line_height
                
                y_position -= 5  # Spațiu între paragrafe
                
                # Pagină nouă dacă e nevoie
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
            
            c.save()
            pdf_buffer.seek(0)
            return pdf_buffer
            
        except Exception as e:
            print(f"Eroare la conversia PDF: {e}")
            return None

    @action(detail=False, methods=['get'], url_path='history')
    def generation_history(self, request):
        """Returnează istoricul documentelor generate."""
        history = GeneratedDocument.objects.select_related(
            'template', 'worker', 'generated_by'
        ).order_by('-generated_at')[:100]
        
        serializer = GeneratedDocumentSerializer(history, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'], url_path='placeholders')
    def list_placeholders(self, request):
        """Listează toate placeholder-ele disponibile pentru template-uri."""
        placeholders = {
            'date_personale': [
                {'key': 'nume', 'description': 'Numele lucrătorului'},
                {'key': 'prenume', 'description': 'Prenumele lucrătorului'},
                {'key': 'nume_complet', 'description': 'Numele și prenumele'},
                {'key': 'cetatenie', 'description': 'Cetățenia'},
                {'key': 'stare_civila', 'description': 'Starea civilă (M/NM)'},
                {'key': 'copii_intretinere', 'description': 'Numărul de copii în întreținere'},
                {'key': 'sex', 'description': 'Sexul (M/F)'},
                {'key': 'data_nasterii', 'description': 'Data nașterii'},
                {'key': 'oras_domiciliu', 'description': 'Orașul de domiciliu'},
            ],
            'pasaport': [
                {'key': 'pasaport_nr', 'description': 'Numărul pașaportului'},
                {'key': 'data_emitere_pass', 'description': 'Data emiterii pașaportului'},
                {'key': 'data_exp_pass', 'description': 'Data expirării pașaportului'},
                {'key': 'autoritate_emitenta_pasaport', 'description': 'Autoritatea emitentă'},
            ],
            'work_permit': [
                {'key': 'dosar_wp_nr', 'description': 'Număr dosar Work Permit'},
                {'key': 'data_solicitare_wp', 'description': 'Data solicitării WP'},
                {'key': 'data_programare_wp', 'description': 'Data programării WP'},
                {'key': 'judet_wp', 'description': 'Județul WP'},
            ],
            'ocupatie': [
                {'key': 'cod_cor', 'description': 'Codul COR'},
                {'key': 'cod_cor_denumire_ro', 'description': 'Denumirea COR în română'},
                {'key': 'cod_cor_denumire_en', 'description': 'Denumirea COR în engleză'},
                {'key': 'functie', 'description': 'Funcția'},
            ],
            'viza': [
                {'key': 'data_solicitare_viza', 'description': 'Data solicitării vizei'},
                {'key': 'data_programare_interviu', 'description': 'Data programării interviu'},
            ],
            'romania': [
                {'key': 'cnp', 'description': 'CNP'},
                {'key': 'data_intrare_ro', 'description': 'Data intrării în România'},
                {'key': 'cim_nr', 'description': 'Număr CIM'},
                {'key': 'data_emitere_cim', 'description': 'Data emiterii CIM'},
                {'key': 'adresa_ro', 'description': 'Adresa din România'},
            ],
            'permis_sedere': [
                {'key': 'data_depunere_ps', 'description': 'Data depunerii PS'},
                {'key': 'data_programare_ps', 'description': 'Data programării PS'},
                {'key': 'data_emitere_ps', 'description': 'Data emiterii PS'},
                {'key': 'data_expirare_ps', 'description': 'Data expirării PS'},
            ],
            'client': [
                {'key': 'client_denumire', 'description': 'Denumirea clientului'},
                {'key': 'client_tara', 'description': 'Țara clientului'},
                {'key': 'client_oras', 'description': 'Orașul clientului'},
                {'key': 'client_judet', 'description': 'Județul clientului'},
                {'key': 'client_adresa', 'description': 'Adresa clientului'},
                {'key': 'client_cod_fiscal', 'description': 'Codul fiscal al clientului'},
            ],
            'altele': [
                {'key': 'status', 'description': 'Status lucrător'},
                {'key': 'observatii', 'description': 'Observații'},
            ],
        }
        return Response(placeholders)