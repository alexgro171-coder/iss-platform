"""
Eco-Fin Views
API endpoints pentru modulul de profitabilitate.
"""
from datetime import datetime
from decimal import Decimal
from django.db.models import Sum, Avg, Count, F, Q
from django.db import transaction
from django.utils import timezone
from django.http import HttpResponse
from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
import openpyxl
from io import BytesIO

# Pentru export PDF
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# Pentru export Word
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from iss.models import Worker, Client, UserProfile, UserRole
from .models import (
    EcoFinSettings, 
    EcoFinImportedRow, 
    EcoFinProcessedRecord, 
    EcoFinImportBatch,
    EcoFinMonthlyReport  # Pentru compatibilitate
)
from .serializers import (
    EcoFinSettingsSerializer,
    EcoFinImportedRowSerializer,
    EcoFinProcessedRecordSerializer,
    EcoFinImportBatchSerializer,
    EcoFinPreviewRowSerializer,
    EcoFinReportSummarySerializer,
    EcoFinMonthlyReportSerializer,  # Pentru compatibilitate
)


class IsManagementOrAdmin(permissions.BasePermission):
    """Permite acces doar pentru Management și Admin."""
    def has_permission(self, request, view):
        if not request.user.is_authenticated:
            return False
        if request.user.is_superuser:
            return True
        try:
            role = request.user.profile.role
        except UserProfile.DoesNotExist:
            return False
        return role in [UserRole.MANAGEMENT, UserRole.ADMIN]


class EcoFinSettingsViewSet(viewsets.ModelViewSet):
    """
    CRUD pentru setările globale Eco-Fin.
    Doar Management/Admin.
    """
    queryset = EcoFinSettings.objects.all()
    serializer_class = EcoFinSettingsSerializer
    permission_classes = [IsManagementOrAdmin]

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    def update(self, request, *args, **kwargs):
        """Verifică dacă setările sunt blocate."""
        instance = self.get_object()
        if instance.is_locked and not request.user.is_superuser:
            return Response(
                {'detail': 'Setările pentru această lună sunt blocate. Contactați Admin.'},
                status=status.HTTP_403_FORBIDDEN
            )
        return super().update(request, *args, **kwargs)

    @action(detail=False, methods=['get'], url_path='current/(?P<year>[0-9]+)/(?P<month>[0-9]+)')
    def get_for_month(self, request, year=None, month=None):
        """Obține setările pentru o lună specifică."""
        try:
            settings = EcoFinSettings.objects.get(year=int(year), month=int(month))
            serializer = self.get_serializer(settings)
            return Response(serializer.data)
        except EcoFinSettings.DoesNotExist:
            return Response({'detail': 'Setări negăsite pentru această lună.'}, status=404)


class EcoFinProcessedRecordViewSet(viewsets.ModelViewSet):
    """
    CRUD pentru înregistrările procesate Eco-Fin.
    Doar Management/Admin pot crea/modifica.
    """
    queryset = EcoFinProcessedRecord.objects.select_related('worker', 'client').all()
    serializer_class = EcoFinProcessedRecordSerializer
    permission_classes = [IsManagementOrAdmin]

    def get_queryset(self):
        """Filtrare după year, month, client_id, worker_id, is_validated."""
        qs = super().get_queryset()
        params = self.request.query_params
        
        year = params.get('year')
        if year:
            qs = qs.filter(year=int(year))
        
        month = params.get('month')
        if month:
            qs = qs.filter(month=int(month))
        
        client_id = params.get('client_id')
        if client_id:
            qs = qs.filter(client_id=int(client_id))
        
        worker_id = params.get('worker_id')
        if worker_id:
            qs = qs.filter(worker_id=int(worker_id))
        
        is_validated = params.get('is_validated')
        if is_validated is not None:
            qs = qs.filter(is_validated=is_validated.lower() == 'true')
        
        return qs.order_by('worker__nume', 'client__denumire')

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    def update(self, request, *args, **kwargs):
        """Verifică dacă raportul e validat - doar Admin poate modifica."""
        instance = self.get_object()
        if instance.is_validated:
            if not request.user.is_superuser:
                return Response(
                    {'detail': 'Înregistrarea este validată. Doar Admin poate modifica.'},
                    status=status.HTTP_403_FORBIDDEN
                )
        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        """Verifică dacă raportul e validat - doar Admin poate șterge."""
        instance = self.get_object()
        if instance.is_validated:
            if not request.user.is_superuser:
                return Response(
                    {'detail': 'Înregistrarea este validată. Doar Admin poate șterge.'},
                    status=status.HTTP_403_FORBIDDEN
                )
        return super().destroy(request, *args, **kwargs)

    @action(detail=False, methods=['get'], url_path='summary')
    def summary(self, request):
        """
        Returnează sumar pentru filtrele specificate.
        GET /api/eco-fin/records/summary/?year=2024&month=12&client_id=1
        """
        qs = self.get_queryset()
        
        if not qs.exists():
            return Response({
                'total_workers': 0,
                'total_hours': 0,
                'total_venit': 0,
                'total_costs': 0,
                'total_profit': 0,
                'average_profit_per_worker': 0,
                'profit_margin_percent': 0,
                'by_client': []
            })
        
        # Calcule agregate
        totals = qs.aggregate(
            total_workers=Count('id'),
            total_hours=Sum('ore_lucrate'),
            total_venit=Sum('venit_generat'),
            total_costs=Sum('cost_salariat_total'),
            total_profit=Sum('profitabilitate'),
        )
        
        # Profit margin
        profit_margin = 0
        if totals['total_venit'] and totals['total_venit'] > 0:
            profit_margin = (totals['total_profit'] / totals['total_venit']) * 100
        
        # Per client (pentru grafic PIE)
        by_client = list(
            qs.values('client__id', 'client__denumire')
            .annotate(
                workers_count=Count('id'),
                total_hours=Sum('ore_lucrate'),
                total_venit=Sum('venit_generat'),
                total_costs=Sum('cost_salariat_total'),
                total_profit=Sum('profitabilitate')
            )
            .order_by('-total_profit')
        )
        
        # Calculăm procentul din profit total pentru fiecare client
        total_profit = totals['total_profit'] or Decimal('0')
        for client_data in by_client:
            client_profit = client_data['total_profit'] or Decimal('0')
            if total_profit > 0:
                client_data['profit_share_percent'] = float((client_profit / total_profit) * 100)
            else:
                client_data['profit_share_percent'] = 0
            
            # Profit margin per client
            if client_data['total_venit'] and client_data['total_venit'] > 0:
                client_data['profit_margin_percent'] = float(
                    (client_data['total_profit'] / client_data['total_venit']) * 100
                )
            else:
                client_data['profit_margin_percent'] = 0
        
        avg_profit = totals['total_profit'] / totals['total_workers'] if totals['total_workers'] else 0
        
        return Response({
            'total_workers': totals['total_workers'] or 0,
            'total_hours': float(totals['total_hours'] or 0),
            'total_venit': float(totals['total_venit'] or 0),
            'total_costs': float(totals['total_costs'] or 0),
            'total_profit': float(totals['total_profit'] or 0),
            'average_profit_per_worker': float(avg_profit),
            'profit_margin_percent': float(profit_margin),
            'by_client': by_client
        })

    @action(detail=False, methods=['post'], url_path='validate-month')
    def validate_month(self, request):
        """
        Validează toate înregistrările dintr-o lună.
        POST /api/eco-fin/records/validate-month/
        Body: { year: 2024, month: 12 }
        """
        year = request.data.get('year')
        month = request.data.get('month')
        
        if not year or not month:
            return Response({'detail': 'Year și month sunt obligatorii.'}, status=400)
        
        year = int(year)
        month = int(month)
        
        # Verifică dacă există înregistrări nevalidate
        records = EcoFinProcessedRecord.objects.filter(
            year=year, month=month, is_validated=False
        )
        
        if not records.exists():
            return Response({
                'detail': f'Nu există înregistrări nevalidate pentru {month:02d}/{year}.'
            }, status=400)
        
        # Validează toate
        now = timezone.now()
        count = records.update(
            is_validated=True,
            validated_at=now,
            validated_by=request.user
        )
        
        # Blochează și setările
        EcoFinSettings.objects.filter(year=year, month=month).update(is_locked=True)
        
        return Response({
            'success': True,
            'validated_count': count,
            'message': f'Au fost validate {count} înregistrări pentru {month:02d}/{year}.'
        })


class EcoFinImportViewSet(viewsets.ViewSet):
    """
    ViewSet pentru import Excel și procesare.
    """
    permission_classes = [IsManagementOrAdmin]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    # Mapare coloane Excel conform specificației
    COLUMN_MAPPING = {
        'nr_cim': ['nr_cim', 'nr cim', 'cim', 'numar_cim', 'numar cim', 'nr.cim'],
        'nume': ['nume', 'name', 'lastname'],
        'prenume': ['prenume', 'firstname', 'first_name'],
        'salariu_brut': ['salariu', 'salariu_brut', 'salariu brut', 'brut', 'salary', 'gross'],
        'ore_lucrate': ['ore', 'ore_lucrate', 'ore lucrate', 'lucrat', 'hours', 'worked_hours'],
        'brut1': ['brut1'],
        'net': ['net', 'salariu_net', 'salariu net', 'net_salary'],
        'retineri': ['retineri', 'rețineri', 'deductions'],
        'rest_plata': ['rest_plata', 'rest plata', 'rest de plata', 'remaining'],
        'cam': ['cam', 'contributie', 'contribuție', 'asigurari', 'asigurări'],
    }

    def _find_column(self, headers, field_name):
        """Găsește coloana pentru un câmp dat."""
        possible_names = self.COLUMN_MAPPING.get(field_name, [field_name])
        for col_idx, header in headers.items():
            header_lower = str(header).lower().strip()
            if header_lower in possible_names:
                return col_idx
        return None

    def _parse_decimal(self, value, default=Decimal('0.00')):
        """Parsează o valoare la Decimal."""
        if value is None:
            return default
        try:
            return Decimal(str(value).replace(',', '.').strip())
        except:
            return default

    @action(detail=False, methods=['post'], url_path='upload')
    def upload(self, request):
        """
        POST /api/eco-fin/import/upload/
        Upload fișier Excel și generare preview cu identificare lucrători.
        
        Coloane așteptate: nr_cim, nume, prenume, salariu (brut), lucrat (ore), 
                          brut1, net, retineri, rest_plata, cam
        """
        file = request.FILES.get('file')
        year = request.data.get('year')
        month = request.data.get('month')
        
        if not file:
            return Response({'detail': 'Fișierul este obligatoriu.'}, status=400)
        if not year or not month:
            return Response({'detail': 'Year și month sunt obligatorii.'}, status=400)
        
        year = int(year)
        month = int(month)
        
        # Verifică dacă luna e deja validată
        existing_validated = EcoFinProcessedRecord.objects.filter(
            year=year, month=month, is_validated=True
        ).exists()
        
        if existing_validated:
            return Response({
                'detail': f'Luna {month:02d}/{year} are deja date validate. Contactați Admin pentru modificări.'
            }, status=400)
        
        # Obține setările pentru lună
        try:
            settings = EcoFinSettings.objects.get(year=year, month=month)
        except EcoFinSettings.DoesNotExist:
            return Response({
                'detail': f'Setările pentru luna {month:02d}/{year} nu sunt configurate. '
                          f'Configurați mai întâi Cheltuielile Indirecte și Costul Concediu.'
            }, status=400)
        
        try:
            # Parsează Excel
            wb = openpyxl.load_workbook(file, data_only=True)
            ws = wb.active
            
            # Găsește header-urile (prima linie)
            headers = {}
            for col_idx, cell in enumerate(ws[1], 1):
                if cell.value:
                    headers[col_idx] = str(cell.value).strip()
            
            # Identifică coloanele
            col_nr_cim = self._find_column(headers, 'nr_cim')
            col_nume = self._find_column(headers, 'nume')
            col_prenume = self._find_column(headers, 'prenume')
            col_salariu = self._find_column(headers, 'salariu_brut')
            col_ore = self._find_column(headers, 'ore_lucrate')
            col_brut1 = self._find_column(headers, 'brut1')
            col_net = self._find_column(headers, 'net')
            col_retineri = self._find_column(headers, 'retineri')
            col_rest = self._find_column(headers, 'rest_plata')
            col_cam = self._find_column(headers, 'cam')
            
            # Validări obligatorii
            missing_cols = []
            if not col_nr_cim:
                missing_cols.append('nr_cim')
            if not col_salariu:
                missing_cols.append('salariu/brut')
            if not col_ore:
                missing_cols.append('ore/lucrat')
            if not col_cam:
                missing_cols.append('cam')
            
            if missing_cols:
                return Response({
                    'detail': f'Coloane obligatorii lipsă: {", ".join(missing_cols)}. '
                              f'Coloane găsite: {list(headers.values())}'
                }, status=400)
            
            # Procesează rândurile
            preview_rows = []
            matched_count = 0
            error_count = 0
            
            for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
                # Citește nr_cim
                nr_cim_val = row[col_nr_cim - 1].value if col_nr_cim else None
                if not nr_cim_val:
                    continue
                
                nr_cim = str(nr_cim_val).strip()
                
                # Citește celelalte valori
                nume = str(row[col_nume - 1].value or '').strip() if col_nume else ''
                prenume = str(row[col_prenume - 1].value or '').strip() if col_prenume else ''
                salariu_brut = self._parse_decimal(row[col_salariu - 1].value if col_salariu else 0)
                ore_lucrate = self._parse_decimal(row[col_ore - 1].value if col_ore else 0)
                brut1 = self._parse_decimal(row[col_brut1 - 1].value if col_brut1 else 0)
                net = self._parse_decimal(row[col_net - 1].value if col_net else 0)
                retineri = self._parse_decimal(row[col_retineri - 1].value if col_retineri else 0)
                rest_plata = self._parse_decimal(row[col_rest - 1].value if col_rest else 0)
                cam = self._parse_decimal(row[col_cam - 1].value if col_cam else 0)
                
                # Caută lucrătorul după nr_CIM
                errors = []
                warnings = []
                worker = Worker.objects.filter(cim_nr__iexact=nr_cim).first()
                
                is_matched = worker is not None
                worker_nume_match = False
                client = None
                
                if not is_matched:
                    errors.append(f'Lucrător cu nr. CIM "{nr_cim}" nu a fost găsit.')
                    error_count += 1
                else:
                    # Verifică concordanța numelui
                    if nume and prenume:
                        worker_nume_lower = (worker.nume or '').lower()
                        worker_prenume_lower = (worker.prenume or '').lower()
                        if worker_nume_lower == nume.lower() and worker_prenume_lower == prenume.lower():
                            worker_nume_match = True
                        else:
                            warnings.append(
                                f'Nume din Excel ({nume} {prenume}) diferă de cel din sistem '
                                f'({worker.nume} {worker.prenume})'
                            )
                    
                    client = worker.client
                    if not client:
                        errors.append('Lucrătorul nu are client asignat.')
                        is_matched = False
                        error_count += 1
                    else:
                        matched_count += 1
                
                # Date din client pentru calcul
                tarif_orar = client.tarif_orar if client else Decimal('0')
                cost_cazare = client.cazare_cost if client else Decimal('0')
                cost_masa = client.masa_cost if client else Decimal('0')
                cost_transport = client.transport_cost if client else Decimal('0')
                
                # Calcule conform formulei
                cost_salarial_complet = salariu_brut + cam
                cota_indirecte = Decimal('0')  # Se calculează după ce știm câți sunt
                cost_concediu = settings.cost_concediu
                
                preview_rows.append({
                    'row_number': row_idx,
                    'nr_cim': nr_cim,
                    'nume': nume,
                    'prenume': prenume,
                    'salariu_brut': float(salariu_brut),
                    'ore_lucrate': float(ore_lucrate),
                    'brut1': float(brut1),
                    'net': float(net),
                    'retineri': float(retineri),
                    'rest_plata': float(rest_plata),
                    'cam': float(cam),
                    'is_matched': is_matched and client is not None,
                    'worker_id': worker.id if worker else None,
                    'worker_nume': worker.nume if worker else None,
                    'worker_prenume': worker.prenume if worker else None,
                    'worker_nume_match': worker_nume_match,
                    'client_id': client.id if client else None,
                    'client_denumire': client.denumire if client else None,
                    'tarif_orar': float(tarif_orar),
                    'cost_cazare': float(cost_cazare),
                    'cost_masa': float(cost_masa),
                    'cost_transport': float(cost_transport),
                    'cost_salarial_complet': float(cost_salarial_complet),
                    'cota_indirecte': 0,  # Se calculează după
                    'cost_concediu': float(cost_concediu),
                    'cost_salariat_total': 0,  # Se calculează după
                    'venit_estimat': 0,  # Se calculează după
                    'profitabilitate_estimata': 0,  # Se calculează după
                    'is_valid': is_matched and client is not None,
                    'errors': errors,
                    'warnings': warnings
                })
            
            # Calculăm cota indirecte per lucrător valid
            if matched_count > 0:
                cota_indirecte_per_worker = float(settings.cheltuieli_indirecte) / matched_count
            else:
                cota_indirecte_per_worker = 0
            
            # Actualizăm calculele pentru rândurile valide
            for row in preview_rows:
                if row['is_valid']:
                    row['cota_indirecte'] = round(cota_indirecte_per_worker, 2)
                    
                    # Cost salariat total
                    cost_salariat_total = (
                        row['cost_salarial_complet'] +
                        row['cost_cazare'] +
                        row['cost_masa'] +
                        row['cost_transport'] +
                        row['cota_indirecte'] +
                        row['cost_concediu']
                    )
                    row['cost_salariat_total'] = round(cost_salariat_total, 2)
                    
                    # Venit estimat
                    venit = row['ore_lucrate'] * row['tarif_orar']
                    row['venit_estimat'] = round(venit, 2)
                    
                    # Profitabilitate
                    row['profitabilitate_estimata'] = round(venit - cost_salariat_total, 2)
            
            # Creăm batch-ul de import
            batch = EcoFinImportBatch.objects.create(
                year=year,
                month=month,
                filename=file.name,
                total_rows=len(preview_rows),
                matched_rows=matched_count,
                error_rows=error_count,
                status=EcoFinImportBatch.Status.PREVIEW,
                imported_by=request.user
            )
            
            return Response({
                'batch_id': batch.id,
                'year': year,
                'month': month,
                'total_rows': len(preview_rows),
                'matched_rows': matched_count,
                'error_rows': error_count,
                'settings': {
                    'cheltuieli_indirecte': float(settings.cheltuieli_indirecte),
                    'cost_concediu': float(settings.cost_concediu),
                    'cota_indirecte_per_worker': cota_indirecte_per_worker
                },
                'preview': preview_rows
            })
            
        except Exception as e:
            import traceback
            return Response({
                'detail': f'Eroare la procesarea fișierului: {str(e)}',
                'traceback': traceback.format_exc()
            }, status=400)

    @action(detail=False, methods=['post'], url_path='process')
    def process_import(self, request):
        """
        POST /api/eco-fin/import/process/
        Procesează și salvează datele din preview în EcoFinProcessedRecord.
        """
        batch_id = request.data.get('batch_id')
        year = request.data.get('year')
        month = request.data.get('month')
        rows = request.data.get('rows', [])
        
        if not year or not month:
            return Response({'detail': 'Year și month sunt obligatorii.'}, status=400)
        
        year = int(year)
        month = int(month)
        
        # Verifică setările
        try:
            settings = EcoFinSettings.objects.get(year=year, month=month)
        except EcoFinSettings.DoesNotExist:
            return Response({'detail': 'Setările pentru această lună nu există.'}, status=400)
        
        # Șterge înregistrările existente nevalidate pentru această lună
        EcoFinProcessedRecord.objects.filter(
            year=year, month=month, is_validated=False
        ).delete()
        
        # Calculăm cota indirecte per lucrător
        valid_rows = [r for r in rows if r.get('is_valid')]
        valid_count = len(valid_rows)
        cota_indirecte = float(settings.cheltuieli_indirecte) / valid_count if valid_count > 0 else 0
        
        created_records = []
        errors = []
        
        with transaction.atomic():
            for row in valid_rows:
                try:
                    worker = Worker.objects.get(id=row['worker_id'])
                    client = Client.objects.get(id=row['client_id'])
                    
                    record = EcoFinProcessedRecord.objects.create(
                        worker=worker,
                        client=client,
                        year=year,
                        month=month,
                        nr_cim=row['nr_cim'],
                        ore_lucrate=Decimal(str(row['ore_lucrate'])),
                        salariu_brut=Decimal(str(row['salariu_brut'])),
                        cam=Decimal(str(row['cam'])),
                        net=Decimal(str(row.get('net', 0))),
                        retineri=Decimal(str(row.get('retineri', 0))),
                        rest_plata=Decimal(str(row.get('rest_plata', 0))),
                        tarif_orar=client.tarif_orar,
                        cost_cazare=client.cazare_cost,
                        cost_masa=client.masa_cost,
                        cost_transport=client.transport_cost,
                        cota_indirecte=Decimal(str(cota_indirecte)),
                        cost_concediu=settings.cost_concediu,
                        is_validated=False,
                        created_by=request.user,
                        notes=row.get('notes', '')
                    )
                    # calculate_costs_and_profit() se apelează automat în save()
                    created_records.append(record.id)
                except Exception as e:
                    errors.append({
                        'row': row.get('row_number'),
                        'nr_cim': row.get('nr_cim'),
                        'error': str(e)
                    })
        
        # Actualizează batch-ul
        if batch_id:
            EcoFinImportBatch.objects.filter(id=batch_id).update(
                status=EcoFinImportBatch.Status.VALIDATED if not errors else EcoFinImportBatch.Status.FAILED,
                processed_rows=len(created_records),
                validated_by=request.user,
                validated_at=timezone.now()
            )
        
        return Response({
            'success': True,
            'created_records': len(created_records),
            'errors': errors,
            'message': f'Au fost create {len(created_records)} înregistrări pentru {month:02d}/{year}.'
        })

    @action(detail=False, methods=['get'], url_path='batches')
    def list_batches(self, request):
        """Lista de batch-uri de import."""
        year = request.query_params.get('year')
        month = request.query_params.get('month')
        
        batches = EcoFinImportBatch.objects.all()
        if year:
            batches = batches.filter(year=int(year))
        if month:
            batches = batches.filter(month=int(month))
        
        batches = batches[:50]
        serializer = EcoFinImportBatchSerializer(batches, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'], url_path='template')
    def download_template(self, request):
        """
        GET /api/eco-fin/import/template/
        Descarcă template Excel pentru import.
        """
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Import Eco-Fin"
        
        # Header-uri conform specificației
        headers = [
            'nr_cim', 'nume', 'prenume', 'salariu', 'lucrat',
            'brut1', 'net', 'retineri', 'rest_plata', 'cam'
        ]
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = openpyxl.styles.Font(bold=True)
        
        # Rând exemplu
        example_row = [
            'CIM001', 'Popescu', 'Ion', 5000.00, 168,
            5000.00, 3200.00, 800.00, 3200.00, 125.00
        ]
        for col, value in enumerate(example_row, 1):
            ws.cell(row=2, column=col, value=value)
        
        # Ajustare lățime coloane
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 15
        
        # Salvare în buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="template_import_ecofin.xlsx"'
        return response


# ==========================================
# RAPOARTE ȘI EXPORT
# ==========================================

@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_by_client(request):
    """
    GET /api/eco-fin/report/client/
    Raport pe Client (lună + an).
    """
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    client_id = request.query_params.get('client_id')
    
    if not year:
        return Response({'detail': 'Year este obligatoriu.'}, status=400)
    
    qs = EcoFinProcessedRecord.objects.filter(year=int(year))
    if month:
        qs = qs.filter(month=int(month))
    if client_id:
        qs = qs.filter(client_id=int(client_id))
    
    # Agregare per client
    clients_data = list(
        qs.values('client__id', 'client__denumire')
        .annotate(
            workers_count=Count('id'),
            total_hours=Sum('ore_lucrate'),
            total_venit=Sum('venit_generat'),
            total_costs=Sum('cost_salariat_total'),
            total_profit=Sum('profitabilitate')
        )
        .order_by('-total_profit')
    )
    
    # Calculăm procente
    total_profit_all = sum(c['total_profit'] or 0 for c in clients_data)
    for c in clients_data:
        c['profit_share_percent'] = (
            float((c['total_profit'] / total_profit_all) * 100) 
            if total_profit_all > 0 else 0
        )
        c['profit_margin_percent'] = (
            float((c['total_profit'] / c['total_venit']) * 100) 
            if c['total_venit'] and c['total_venit'] > 0 else 0
        )
    
    return Response({
        'year': int(year),
        'month': int(month) if month else None,
        'clients': clients_data,
        'totals': {
            'total_clients': len(clients_data),
            'total_profit': float(total_profit_all)
        }
    })


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_workers_by_client(request):
    """
    GET /api/eco-fin/report/workers/
    Raport lucrători la un Client.
    """
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    client_id = request.query_params.get('client_id')
    
    if not year or not client_id:
        return Response({'detail': 'Year și client_id sunt obligatorii.'}, status=400)
    
    qs = EcoFinProcessedRecord.objects.filter(
        year=int(year),
        client_id=int(client_id)
    ).select_related('worker', 'client')
    
    if month:
        qs = qs.filter(month=int(month))
    
    records = qs.order_by('worker__nume')
    serializer = EcoFinProcessedRecordSerializer(records, many=True)
    
    # Totaluri
    totals = qs.aggregate(
        total_workers=Count('id'),
        total_hours=Sum('ore_lucrate'),
        total_venit=Sum('venit_generat'),
        total_costs=Sum('cost_salariat_total'),
        total_profit=Sum('profitabilitate')
    )
    
    return Response({
        'year': int(year),
        'month': int(month) if month else None,
        'client_id': int(client_id),
        'workers': serializer.data,
        'totals': {
            'total_workers': totals['total_workers'] or 0,
            'total_hours': float(totals['total_hours'] or 0),
            'total_venit': float(totals['total_venit'] or 0),
            'total_costs': float(totals['total_costs'] or 0),
            'total_profit': float(totals['total_profit'] or 0)
        }
    })


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_all_clients(request):
    """
    GET /api/eco-fin/report/all/
    Raport toți clienții (lună selectată) - pentru grafic PIE.
    """
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    
    if not year or not month:
        return Response({'detail': 'Year și month sunt obligatorii.'}, status=400)
    
    qs = EcoFinProcessedRecord.objects.filter(
        year=int(year),
        month=int(month)
    )
    
    # Agregare per client
    clients_data = list(
        qs.values('client__id', 'client__denumire')
        .annotate(
            workers_count=Count('id'),
            total_hours=Sum('ore_lucrate'),
            tarif_orar_avg=Avg('tarif_orar'),
            total_venit=Sum('venit_generat'),
            total_costs=Sum('cost_salariat_total'),
            total_profit=Sum('profitabilitate')
        )
        .order_by('-total_profit')
    )
    
    # Calculăm procente pentru grafic PIE
    total_profit_all = sum(c['total_profit'] or 0 for c in clients_data)
    for c in clients_data:
        c['profit_share_percent'] = (
            float((c['total_profit'] / total_profit_all) * 100) 
            if total_profit_all > 0 else 0
        )
        # Valoare de facturat
        c['valoare_facturare'] = float(c['total_hours'] or 0) * float(c['tarif_orar_avg'] or 0)
    
    return Response({
        'year': int(year),
        'month': int(month),
        'clients': clients_data,
        'chart_data': [
            {
                'name': c['client__denumire'],
                'value': float(c['total_profit'] or 0),
                'percent': c['profit_share_percent']
            }
            for c in clients_data if c['total_profit'] and c['total_profit'] > 0
        ],
        'totals': {
            'total_clients': len(clients_data),
            'total_workers': sum(c['workers_count'] for c in clients_data),
            'total_hours': sum(float(c['total_hours'] or 0) for c in clients_data),
            'total_venit': sum(float(c['total_venit'] or 0) for c in clients_data),
            'total_costs': sum(float(c['total_costs'] or 0) for c in clients_data),
            'total_profit': float(total_profit_all)
        }
    })


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_interval(request):
    """
    GET /api/eco-fin/report/interval/
    Rapoarte multi-lunare / multi-an.
    """
    year_start = request.query_params.get('year_start')
    month_start = request.query_params.get('month_start')
    year_end = request.query_params.get('year_end')
    month_end = request.query_params.get('month_end')
    client_id = request.query_params.get('client_id')
    group_by = request.query_params.get('group_by', 'month')  # 'month' sau 'client'
    
    if not year_start or not year_end:
        return Response({'detail': 'year_start și year_end sunt obligatorii.'}, status=400)
    
    # Construiește filtrele
    qs = EcoFinProcessedRecord.objects.all()
    
    # Filtru interval
    if month_start:
        qs = qs.filter(
            Q(year__gt=int(year_start)) | 
            Q(year=int(year_start), month__gte=int(month_start))
        )
    else:
        qs = qs.filter(year__gte=int(year_start))
    
    if month_end:
        qs = qs.filter(
            Q(year__lt=int(year_end)) | 
            Q(year=int(year_end), month__lte=int(month_end))
        )
    else:
        qs = qs.filter(year__lte=int(year_end))
    
    if client_id:
        qs = qs.filter(client_id=int(client_id))
    
    # Agregare
    if group_by == 'month':
        data = list(
            qs.values('year', 'month')
            .annotate(
                workers_count=Count('id'),
                total_hours=Sum('ore_lucrate'),
                total_venit=Sum('venit_generat'),
                total_costs=Sum('cost_salariat_total'),
                total_profit=Sum('profitabilitate')
            )
            .order_by('year', 'month')
        )
    else:  # group_by == 'client'
        data = list(
            qs.values('client__id', 'client__denumire')
            .annotate(
                workers_count=Count('id'),
                total_hours=Sum('ore_lucrate'),
                total_venit=Sum('venit_generat'),
                total_costs=Sum('cost_salariat_total'),
                total_profit=Sum('profitabilitate')
            )
            .order_by('-total_profit')
        )
    
    # Totaluri generale
    totals = qs.aggregate(
        total_records=Count('id'),
        total_hours=Sum('ore_lucrate'),
        total_venit=Sum('venit_generat'),
        total_costs=Sum('cost_salariat_total'),
        total_profit=Sum('profitabilitate')
    )
    
    return Response({
        'period': {
            'year_start': int(year_start),
            'month_start': int(month_start) if month_start else 1,
            'year_end': int(year_end),
            'month_end': int(month_end) if month_end else 12
        },
        'group_by': group_by,
        'data': data,
        'totals': {
            'total_records': totals['total_records'] or 0,
            'total_hours': float(totals['total_hours'] or 0),
            'total_venit': float(totals['total_venit'] or 0),
            'total_costs': float(totals['total_costs'] or 0),
            'total_profit': float(totals['total_profit'] or 0)
        }
    })


# ==========================================
# RAPOARTE FINANCIARE (REST PLATĂ, REȚINERI)
# ==========================================

@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_rest_plata_total(request):
    """
    GET /api/eco-fin/report/rest-plata/
    Raport total rest de plată către lucrătorii activi.
    
    Parametri:
    - year (obligatoriu)
    - month (opțional - dacă nu e specificat, toate lunile din an)
    """
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    
    if not year:
        return Response({'detail': 'Year este obligatoriu.'}, status=400)
    
    qs = EcoFinProcessedRecord.objects.filter(year=int(year))
    if month:
        qs = qs.filter(month=int(month))
    
    # Totaluri generale
    totals = qs.aggregate(
        total_workers=Count('id'),
        total_rest_plata=Sum('rest_plata'),
        total_retineri=Sum('retineri'),
        total_net=Sum('net'),
        total_brut=Sum('salariu_brut'),
        total_cam=Sum('cam')
    )
    
    # Per lună (pentru grafic)
    per_month = list(
        qs.values('year', 'month')
        .annotate(
            workers_count=Count('id'),
            rest_plata=Sum('rest_plata'),
            retineri=Sum('retineri'),
            net=Sum('net')
        )
        .order_by('year', 'month')
    )
    
    # Detalii per lucrător (primele 100)
    workers_detail = list(
        qs.values(
            'worker__id', 'worker__nume', 'worker__prenume', 
            'worker__pasaport_nr', 'nr_cim'
        )
        .annotate(
            total_rest_plata=Sum('rest_plata'),
            total_retineri=Sum('retineri'),
            total_net=Sum('net'),
            months_count=Count('id')
        )
        .order_by('-total_rest_plata')[:100]
    )
    
    return Response({
        'year': int(year),
        'month': int(month) if month else None,
        'totals': {
            'total_workers': totals['total_workers'] or 0,
            'total_rest_plata': float(totals['total_rest_plata'] or 0),
            'total_retineri': float(totals['total_retineri'] or 0),
            'total_net': float(totals['total_net'] or 0),
            'total_brut': float(totals['total_brut'] or 0),
            'total_cam': float(totals['total_cam'] or 0)
        },
        'per_month': per_month,
        'workers': workers_detail
    })


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_rest_plata_by_client(request):
    """
    GET /api/eco-fin/report/rest-plata-client/
    Raport rest de plată către lucrătorii activi grupat pe client.
    
    Parametri:
    - year (obligatoriu)
    - month (opțional)
    - client_id (opțional - pentru detalii pe un singur client)
    """
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    client_id = request.query_params.get('client_id')
    
    if not year:
        return Response({'detail': 'Year este obligatoriu.'}, status=400)
    
    qs = EcoFinProcessedRecord.objects.filter(year=int(year))
    if month:
        qs = qs.filter(month=int(month))
    if client_id:
        qs = qs.filter(client_id=int(client_id))
    
    # Per client
    clients_data = list(
        qs.values('client__id', 'client__denumire')
        .annotate(
            workers_count=Count('id'),
            total_rest_plata=Sum('rest_plata'),
            total_retineri=Sum('retineri'),
            total_net=Sum('net'),
            total_brut=Sum('salariu_brut')
        )
        .order_by('-total_rest_plata')
    )
    
    # Totaluri
    total_rest_plata_all = sum(c['total_rest_plata'] or 0 for c in clients_data)
    
    # Calculăm procentele
    for c in clients_data:
        c['rest_plata_share_percent'] = (
            float((c['total_rest_plata'] / total_rest_plata_all) * 100)
            if total_rest_plata_all > 0 else 0
        )
    
    # Dacă avem client_id, adăugăm detalii per lucrător
    workers_detail = []
    if client_id:
        workers_detail = list(
            qs.values(
                'worker__id', 'worker__nume', 'worker__prenume',
                'worker__pasaport_nr', 'nr_cim', 'month'
            )
            .annotate(
                rest_plata=Sum('rest_plata'),
                retineri=Sum('retineri'),
                net=Sum('net')
            )
            .order_by('worker__nume', 'month')
        )
    
    return Response({
        'year': int(year),
        'month': int(month) if month else None,
        'client_id': int(client_id) if client_id else None,
        'clients': clients_data,
        'workers_detail': workers_detail,
        'totals': {
            'total_clients': len(clients_data),
            'total_rest_plata': float(total_rest_plata_all),
            'total_retineri': sum(float(c['total_retineri'] or 0) for c in clients_data),
            'total_net': sum(float(c['total_net'] or 0) for c in clients_data)
        },
        'chart_data': [
            {
                'name': c['client__denumire'],
                'value': float(c['total_rest_plata'] or 0),
                'percent': c['rest_plata_share_percent']
            }
            for c in clients_data if c['total_rest_plata'] and c['total_rest_plata'] > 0
        ]
    })


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_retineri(request):
    """
    GET /api/eco-fin/report/retineri/
    Raport detaliat rețineri.
    
    Parametri:
    - year (obligatoriu)
    - month (opțional)
    - client_id (opțional)
    """
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    client_id = request.query_params.get('client_id')
    
    if not year:
        return Response({'detail': 'Year este obligatoriu.'}, status=400)
    
    qs = EcoFinProcessedRecord.objects.filter(year=int(year))
    if month:
        qs = qs.filter(month=int(month))
    if client_id:
        qs = qs.filter(client_id=int(client_id))
    
    # Totaluri
    totals = qs.aggregate(
        total_workers=Count('id'),
        total_retineri=Sum('retineri'),
        total_brut=Sum('salariu_brut'),
        total_net=Sum('net'),
        total_rest_plata=Sum('rest_plata')
    )
    
    # Per client
    per_client = list(
        qs.values('client__id', 'client__denumire')
        .annotate(
            workers_count=Count('id'),
            total_retineri=Sum('retineri'),
            total_brut=Sum('salariu_brut'),
            avg_retineri=Avg('retineri')
        )
        .order_by('-total_retineri')
    )
    
    # Per lună (pentru trend)
    per_month = list(
        qs.values('year', 'month')
        .annotate(
            workers_count=Count('id'),
            total_retineri=Sum('retineri'),
            avg_retineri=Avg('retineri')
        )
        .order_by('year', 'month')
    )
    
    # Procent rețineri din brut
    retineri_percent = 0
    if totals['total_brut'] and totals['total_brut'] > 0:
        retineri_percent = (totals['total_retineri'] / totals['total_brut']) * 100
    
    return Response({
        'year': int(year),
        'month': int(month) if month else None,
        'client_id': int(client_id) if client_id else None,
        'totals': {
            'total_workers': totals['total_workers'] or 0,
            'total_retineri': float(totals['total_retineri'] or 0),
            'total_brut': float(totals['total_brut'] or 0),
            'total_net': float(totals['total_net'] or 0),
            'total_rest_plata': float(totals['total_rest_plata'] or 0),
            'retineri_percent_of_brut': float(retineri_percent)
        },
        'per_client': per_client,
        'per_month': per_month
    })


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def report_financial_summary(request):
    """
    GET /api/eco-fin/report/financial-summary/
    Raport sumar financiar complet.
    
    Parametri:
    - year (obligatoriu)
    - month (opțional)
    """
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    
    if not year:
        return Response({'detail': 'Year este obligatoriu.'}, status=400)
    
    qs = EcoFinProcessedRecord.objects.filter(year=int(year))
    if month:
        qs = qs.filter(month=int(month))
    
    # Totaluri complete
    totals = qs.aggregate(
        total_workers=Count('id'),
        total_ore=Sum('ore_lucrate'),
        total_brut=Sum('salariu_brut'),
        total_cam=Sum('cam'),
        total_cost_salarial=Sum('cost_salarial_complet'),
        total_net=Sum('net'),
        total_retineri=Sum('retineri'),
        total_rest_plata=Sum('rest_plata'),
        total_cazare=Sum('cost_cazare'),
        total_masa=Sum('cost_masa'),
        total_transport=Sum('cost_transport'),
        total_indirecte=Sum('cota_indirecte'),
        total_concediu=Sum('cost_concediu'),
        total_cost_total=Sum('cost_salariat_total'),
        total_venit=Sum('venit_generat'),
        total_profit=Sum('profitabilitate')
    )
    
    # Calculăm ratii și procente
    venit = float(totals['total_venit'] or 0)
    cost = float(totals['total_cost_total'] or 0)
    profit = float(totals['total_profit'] or 0)
    brut = float(totals['total_brut'] or 0)
    
    return Response({
        'year': int(year),
        'month': int(month) if month else None,
        'summary': {
            # Venituri și Profit
            'total_workers': totals['total_workers'] or 0,
            'total_ore': float(totals['total_ore'] or 0),
            'total_venit': venit,
            'total_profit': profit,
            'profit_margin': (profit / venit * 100) if venit > 0 else 0,
            
            # Costuri salariale
            'salarii': {
                'brut': brut,
                'cam': float(totals['total_cam'] or 0),
                'cost_salarial_complet': float(totals['total_cost_salarial'] or 0),
                'net': float(totals['total_net'] or 0),
                'retineri': float(totals['total_retineri'] or 0),
                'rest_plata': float(totals['total_rest_plata'] or 0),
                'retineri_percent': (float(totals['total_retineri'] or 0) / brut * 100) if brut > 0 else 0
            },
            
            # Alte costuri
            'alte_costuri': {
                'cazare': float(totals['total_cazare'] or 0),
                'masa': float(totals['total_masa'] or 0),
                'transport': float(totals['total_transport'] or 0),
                'indirecte': float(totals['total_indirecte'] or 0),
                'concediu': float(totals['total_concediu'] or 0)
            },
            
            # Total costuri
            'total_costuri': cost,
            
            # Ratii
            'ratii': {
                'cost_per_ora': (cost / float(totals['total_ore'])) if totals['total_ore'] else 0,
                'venit_per_ora': (venit / float(totals['total_ore'])) if totals['total_ore'] else 0,
                'profit_per_lucrator': (profit / totals['total_workers']) if totals['total_workers'] else 0,
                'cost_salarial_percent': (float(totals['total_cost_salarial'] or 0) / cost * 100) if cost > 0 else 0
            }
        }
    })


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def export_pdf(request):
    """
    GET /api/eco-fin/export/pdf/
    Export raport în format PDF.
    """
    if not HAS_REPORTLAB:
        return Response({'detail': 'Biblioteca reportlab nu este instalată.'}, status=500)
    
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    client_id = request.query_params.get('client_id')
    report_type = request.query_params.get('type', 'summary')  # 'summary', 'detailed', 'client'
    
    if not year:
        return Response({'detail': 'Year este obligatoriu.'}, status=400)
    
    # Obține datele
    qs = EcoFinProcessedRecord.objects.filter(year=int(year))
    if month:
        qs = qs.filter(month=int(month))
    if client_id:
        qs = qs.filter(client_id=int(client_id))
    
    qs = qs.select_related('worker', 'client').order_by('client__denumire', 'worker__nume')
    
    # Creează PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=landscape(A4),
        rightMargin=1*cm, leftMargin=1*cm,
        topMargin=1*cm, bottomMargin=1*cm
    )
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=1  # Center
    )
    
    elements = []
    
    # Titlu
    period = f"{month:02d}/{year}" if month else str(year)
    elements.append(Paragraph(f"Raport Eco-Fin - {period}", title_style))
    elements.append(Spacer(1, 20))
    
    # Tabel
    if report_type == 'detailed':
        headers = ['Lucrător', 'Client', 'Ore', 'Salariu', 'CAM', 'Cost Total', 'Venit', 'Profit']
        data = [headers]
        
        for record in qs:
            data.append([
                f"{record.worker.nume} {record.worker.prenume}",
                record.client.denumire,
                f"{record.ore_lucrate:.1f}",
                f"{record.salariu_brut:.2f}",
                f"{record.cam:.2f}",
                f"{record.cost_salariat_total:.2f}",
                f"{record.venit_generat:.2f}",
                f"{record.profitabilitate:.2f}"
            ])
    else:
        # Summary per client
        clients_agg = qs.values('client__denumire').annotate(
            workers=Count('id'),
            hours=Sum('ore_lucrate'),
            costs=Sum('cost_salariat_total'),
            revenue=Sum('venit_generat'),
            profit=Sum('profitabilitate')
        ).order_by('-profit')
        
        headers = ['Client', 'Lucrători', 'Ore', 'Costuri', 'Venituri', 'Profit']
        data = [headers]
        
        for c in clients_agg:
            data.append([
                c['client__denumire'],
                str(c['workers']),
                f"{c['hours']:.1f}",
                f"{c['costs']:.2f}",
                f"{c['revenue']:.2f}",
                f"{c['profit']:.2f}"
            ])
    
    # Creare tabel
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f1f5f9')])
    ]))
    
    elements.append(table)
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="raport_ecofin_{period}.pdf"'
    return response


@api_view(['GET'])
@permission_classes([IsManagementOrAdmin])
def export_word(request):
    """
    GET /api/eco-fin/export/word/
    Export raport în format Word.
    """
    if not HAS_DOCX:
        return Response({'detail': 'Biblioteca python-docx nu este instalată.'}, status=500)
    
    year = request.query_params.get('year')
    month = request.query_params.get('month')
    client_id = request.query_params.get('client_id')
    
    if not year:
        return Response({'detail': 'Year este obligatoriu.'}, status=400)
    
    # Obține datele
    qs = EcoFinProcessedRecord.objects.filter(year=int(year))
    if month:
        qs = qs.filter(month=int(month))
    if client_id:
        qs = qs.filter(client_id=int(client_id))
    
    # Agregare per client
    clients_agg = qs.values('client__denumire').annotate(
        workers=Count('id'),
        hours=Sum('ore_lucrate'),
        costs=Sum('cost_salariat_total'),
        revenue=Sum('venit_generat'),
        profit=Sum('profitabilitate')
    ).order_by('-profit')
    
    # Creează document Word
    doc = Document()
    
    # Titlu
    period = f"{month:02d}/{year}" if month else str(year)
    doc.add_heading(f'Raport Eco-Fin - {period}', 0)
    
    # Tabel
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    headers = ['Client', 'Lucrători', 'Ore', 'Costuri', 'Venituri', 'Profit']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Date
    for c in clients_agg:
        row = table.add_row().cells
        row[0].text = c['client__denumire']
        row[1].text = str(c['workers'])
        row[2].text = f"{c['hours']:.1f}"
        row[3].text = f"{c['costs']:.2f} RON"
        row[4].text = f"{c['revenue']:.2f} RON"
        row[5].text = f"{c['profit']:.2f} RON"
    
    # Salvare în buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="raport_ecofin_{period}.docx"'
    return response


# ==========================================
# COMPATIBILITATE CU VECHIUL MODEL
# ==========================================

class EcoFinMonthlyReportViewSet(viewsets.ModelViewSet):
    """
    [DEPRECIAT] ViewSet pentru rapoartele lunare vechi.
    Folosiți EcoFinProcessedRecordViewSet în schimb.
    """
    queryset = EcoFinMonthlyReport.objects.select_related('worker', 'client').all()
    serializer_class = EcoFinMonthlyReportSerializer
    permission_classes = [IsManagementOrAdmin]

    def get_queryset(self):
        qs = super().get_queryset()
        params = self.request.query_params
        
        year = params.get('year')
        if year:
            qs = qs.filter(year=int(year))
        month = params.get('month')
        if month:
            qs = qs.filter(month=int(month))
        client_id = params.get('client_id')
        if client_id:
            qs = qs.filter(client_id=int(client_id))
        
        return qs.order_by('worker__nume', 'client__denumire')

    @action(detail=False, methods=['get'], url_path='summary')
    def summary(self, request):
        """Sumar pentru compatibilitate."""
        qs = self.get_queryset()
        
        if not qs.exists():
            return Response({
                'total_workers': 0,
                'total_hours': 0,
                'total_salary_cost': 0,
                'total_revenue': 0,
                'total_costs': 0,
                'total_profit': 0,
                'average_profit_per_worker': 0,
                'by_client': []
            })
        
        totals = qs.aggregate(
            total_workers=Count('id'),
            total_hours=Sum('hours_worked'),
            total_salary_cost=Sum('salary_cost'),
            total_profit=Sum('profit_brut'),
        )
        
        total_revenue = sum(r.hours_worked * r.tarif_orar for r in qs)
        total_costs = sum(
            r.salary_cost + r.cost_cazare + r.cost_masa + 
            r.cost_transport + r.cost_concediu + r.cheltuieli_indirecte
            for r in qs
        )
        
        by_client = list(
            qs.values('client__id', 'client__denumire')
            .annotate(
                workers_count=Count('id'),
                total_hours=Sum('hours_worked'),
                total_profit=Sum('profit_brut')
            )
            .order_by('-total_profit')
        )
        
        avg_profit = totals['total_profit'] / totals['total_workers'] if totals['total_workers'] else 0
        
        return Response({
            'total_workers': totals['total_workers'] or 0,
            'total_hours': float(totals['total_hours'] or 0),
            'total_salary_cost': float(totals['total_salary_cost'] or 0),
            'total_revenue': float(total_revenue),
            'total_costs': float(total_costs),
            'total_profit': float(totals['total_profit'] or 0),
            'average_profit_per_worker': float(avg_profit),
            'by_client': by_client
        })
